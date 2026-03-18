import re
import pandas as pd
from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Cm, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_ORIENTATION
from docx2pdf import convert


class DocumentConverter:
    @staticmethod
    def clean_text(text):
        text = str(text)
        return re.sub(r'[\x00-\x08\x0b\x0c\x0e-\x1f\x7f-\x9f]', '', text)

    @staticmethod
    def add_page_number(run):
        fldChar1 = OxmlElement('w:fldChar')
        fldChar1.set(qn('w:fldCharType'), 'begin')
        instrText = OxmlElement('w:instrText')
        instrText.set(qn('xml:space'), 'preserve')
        instrText.text = "PAGE"
        fldChar2 = OxmlElement('w:fldChar')
        fldChar2.set(qn('w:fldCharType'), 'end')

        run._r.append(fldChar1)
        run._r.append(instrText)
        run._r.append(fldChar2)

    @staticmethod
    def generate_docx(excel_path, output_path, title, config):
        df = pd.read_excel(excel_path)
        df = df.fillna("")
        doc = Document()

        # Odczytujemy konfigurację dokumentu
        font_name = config.get("font_family", "Arial")
        font_size = int(config.get("font_size", 12))
        line_spacing = float(config.get("line_spacing", 1.15))
        margin = Cm(float(str(config.get("margin", 2.0)).replace(',', '.')))
        orientation = config.get("orientation", "Horizontal")
        is_bold = config.get("bold_headers", True)
        is_italic = config.get("italic_headers", False)
        table_style = config.get("table_style", "Light Grid")
        alignment_str = config.get("alignment", "To Left")
        align_map = {
            "To Left": WD_ALIGN_PARAGRAPH.LEFT,
            "Centered": WD_ALIGN_PARAGRAPH.CENTER,
            "To Right": WD_ALIGN_PARAGRAPH.RIGHT,
            "Justify": WD_ALIGN_PARAGRAPH.JUSTIFY
        }
        text_alignment = align_map.get(alignment_str, WD_ALIGN_PARAGRAPH.LEFT)

        style = doc.styles['Normal']
        style.font.name = font_name
        style.font.size = Pt(font_size)
        style._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)

        # Marginesy oraz numeracja stron (stopka)
        for section in doc.sections:
            section.top_margin = section.bottom_margin = margin
            section.left_margin = section.right_margin = margin
            if orientation == "Horizontal":
                section.orientation = WD_ORIENTATION.LANDSCAPE
                new_width, new_height = section.page_height, section.page_width
                section.page_width = new_width
                section.page_height = new_height
            else:
                section.orientation = WD_ORIENTATION.PORTRAIT

            if config.get("page_numbers", True):
                footer = section.footer
                footer_p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
                footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                footer_run = footer_p.add_run("Strona ")
                footer_run.font.name = font_name
                DocumentConverter.add_page_number(footer_p.add_run())

        # Tytuł raportu
        if title:
            heading = doc.add_heading(title, 0)
            heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in heading.runs:
                run.font.name = font_name

        if config.get("format_type") == "Table":
            table = doc.add_table(rows=1, cols=len(df.columns))
            table.style = table_style
            # Nagłówki kolumny
            hdr_cells = table.rows[0].cells
            for i, column in enumerate(df.columns):
                hdr_cells[i].text = str(column)
                for paragraph in hdr_cells[i].paragraphs:
                    paragraph.alignment = text_alignment
                    paragraph.paragraph_format.line_spacing = line_spacing
                    for run in paragraph.runs:
                        run.font.name = font_name
                        run.font.size = Pt(font_size)
                        run.font.bold = is_bold
                        run.font.italic = is_italic

            # Powtarzanie nagłówków co strona
            tr = table.rows[0]._tr
            trPr = tr.get_or_add_trPr()
            tblHeader = OxmlElement('w:tblHeader')
            tblHeader.set(qn('w:val'), 'true')
            trPr.append(tblHeader)
            # Zapobieganiu łamaniu się komórek na strony
            cantSplit = OxmlElement('w:cantSplit')
            cantSplit.set(qn('w:val'), 'true')
            trPr.append(cantSplit)
            # Komórki z danymi
            for _, row in df.iterrows():
                row_cells = table.add_row().cells
                for i, val in enumerate(row):
                    row_cells[i].text = DocumentConverter.clean_text(val)
                    for paragraph in row_cells[i].paragraphs:
                        paragraph.alignment = text_alignment
                        paragraph.paragraph_format.line_spacing = line_spacing
                        for run in paragraph.runs:
                            run.font.name = font_name
                            run.font.size = Pt(font_size)
        else:
            # Lista
            columns_to_ignore = ['lp', 'l.p.', 'l.p', 'lp.', 'id', 'no', 'nr']
            for index, row in df.iterrows():
                # Separatory z Numeracją rekordów
                separator = doc.add_paragraph()
                separator.paragraph_format.keep_with_next = True
                separator.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run_sep = separator.add_run(f"Rekord {index + 1}")
                run_sep.bold = True
                run_sep.font.name = font_name
                run_sep.font.size = Pt(font_size + 2)
                # Właściwe wiersze z danymi
                items = list(row.items())
                valid_items = [(name, val) for name, val in items if str(name).strip().lower() not in columns_to_ignore]
                for i, (name, val) in enumerate(valid_items):
                    p = doc.add_paragraph()
                    p.alignment = text_alignment
                    p.paragraph_format.line_spacing = line_spacing
                    if i < len(items) - 1:
                        p.paragraph_format.keep_with_next = True
                    # Klucz dla elementu listy
                    run_col = p.add_run(f"{DocumentConverter.clean_text(name)}: ")
                    run_col.font.bold = is_bold
                    run_col.font.italic = is_italic
                    run_col.font.name = font_name
                    run_col.font.size = Pt(font_size)
                    # Wartość dla elementu listy
                    run_val = p.add_run(str(DocumentConverter.clean_text(val)))
                    run_val.font.name = font_name
                    run_val.font.size = Pt(font_size)

        doc.save(output_path)

    @staticmethod
    def generate_pdf(docx_path, output_path):
        convert(docx_path, output_path)
